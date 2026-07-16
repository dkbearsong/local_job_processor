import re
from typing import Any, Optional
from docx import Document
from docx.shared import Pt, Inches
import docx.opc.constants

def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Adds a clickable hyperlink to a paragraph."""
    from docx.oxml.parser import OxmlElement
    from docx.oxml.ns import qn
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def get_first_url(p: Any) -> Optional[str]:
    """Retrieves the first URL found inside a paragraph or its text representation."""
    from docx.oxml.ns import qn
    try:
        hyperlinks = p._p.xpath('.//w:hyperlink')
        for hl in hyperlinks:
            rId = hl.get(qn('r:id'))
            if rId and rId in p.part.rels:
                return p.part.rels[rId]._target
    except Exception:
        pass
            
    text = p.text.strip()
    match = re.search(r'(https?://\S+)', text)
    if match:
        return match.group(1)
        
    return None

def format_resume_docx(docx_path: str) -> None:
    """
    Format the generated .docx resume with customized typography, layout margins, 
    bullet indentation, and inline hyperlinks.
    """
    doc = Document(docx_path)
    
    style = doc.styles['Normal']
    font_name = getattr(style.font, 'name', None)  # type: ignore[attr-defined]
    if font_name is None or 'Calibri' not in font_name:
        style.font.name = 'Calibri'  # type: ignore[attr-defined]
    style.font.size = Pt(11)  # type: ignore[attr-defined]
    style.paragraph_format.line_spacing = 1.15  # type: ignore[attr-defined]
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    in_projects = False
    in_skills = False
    paragraphs_to_delete = []
    current_project_title_p = None
    current_skill_category_p = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        is_main_heading = False
        if p.style is not None and p.style.name is not None and p.style.name.startswith('Heading'):  # type: ignore[attr-defined]
            lower_text = text.lower()
            if lower_text in ['summary', 'experience', 'projects', 'skills', 'education', 'contact information']:
                p.style = doc.styles['Heading 2']  # type: ignore[assignment]
                in_projects = (lower_text == 'projects')
                in_skills = (lower_text == 'skills')
                is_main_heading = True
        
        if is_main_heading:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            continue
            
        if in_projects:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            url = get_first_url(p)
            is_link_line = False
            if url:
                clean_text_for_link = text
                if clean_text_for_link.startswith('- ') or clean_text_for_link.startswith('* ') or clean_text_for_link.startswith('• '):
                    clean_text_for_link = clean_text_for_link[2:].strip()
                text_without_url = clean_text_for_link.replace(url, '').strip()
                if not text_without_url or text_without_url.lower() in ['link:', 'github:', 'source:', 'repo:', 'repository:', 'project link:', 'view project:', 'link']:
                    is_link_line = True
                elif len(text_without_url.split()) <= 3:
                    is_link_line = True
            
            if (p.style is not None and p.style.name == 'Heading 3') or (p.runs and p.runs[0].bold and not text.lower().startswith('tech stack')):  # type: ignore[attr-defined]
                p.style = doc.styles['Normal']  # type: ignore[assignment]
                p.paragraph_format.space_before = Pt(12) # Line space above project listing
                for r in p.runs:
                    r.bold = True
                current_project_title_p = p
                
                # Check if the title paragraph itself contains a URL (e.g. "### Project Name (https://...)")
                if url and not is_link_line:
                    # The title has a URL embedded in it - strip the URL from the title text
                    title_text = text_without_url if url and 'text_without_url' in locals() and text_without_url else text
                    # Clean up any leftover parentheses/brackets around where the URL was
                    title_text = title_text.rstrip('(').rstrip('[').strip()
                    print(f"[HYPERLINK] Applying URL '{url}' to project title: '{title_text}'")
                    p.clear()
                    if url is not None:
                        add_hyperlink(p, title_text, url)
                    if p.runs:
                        p.runs[-1].bold = True
                        
            elif is_link_line:
                if current_project_title_p:
                    if not current_project_title_p._p.xpath('.//w:hyperlink'):
                        title_text = current_project_title_p.text
                        print(f"[HYPERLINK] Applying URL '{url}' to project title: '{title_text}'")
                        current_project_title_p.clear()
                        if url is not None:
                            add_hyperlink(current_project_title_p, title_text, url)
                        if current_project_title_p.runs:
                            current_project_title_p.runs[-1].bold = True
                paragraphs_to_delete.append(p)
            elif text.lower().startswith('tech stack'):
                p.style = doc.styles['Normal']  # type: ignore[assignment]
                p.paragraph_format.left_indent = Inches(0.25)
                p.clear()
                run = p.add_run(text)
                run.italic = True
                if current_project_title_p is not None:
                    current_project_title_p._element.addnext(p._element)
            else:
                p.style = doc.styles['Normal']  # type: ignore[assignment]
                p.paragraph_format.left_indent = Inches(0.25)
                clean_text = text
                if clean_text.startswith('- ') or clean_text.startswith('* ') or clean_text.startswith('• '):
                    clean_text = clean_text[2:]
                p.clear()
                p.add_run(clean_text)

        elif in_skills:
            p.style = doc.styles['Normal']  # type: ignore[assignment]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            
            clean_text = text
            if clean_text.startswith('- ') or clean_text.startswith('* ') or clean_text.startswith('• '):
                clean_text = clean_text[2:]
            
            if ':' in clean_text:
                title, rest = clean_text.split(':', 1)
                p.clear()
                r1 = p.add_run(title.strip() + ": ")
                r1.bold = True
                if rest.strip():
                    p.add_run(rest.strip())
                current_skill_category_p = p
            elif p.style is not None and p.style.name == 'Heading 3' or (p.runs and p.runs[0].bold):  # type: ignore[attr-defined]
                p.clear()
                r1 = p.add_run(clean_text.strip() + ": ")
                r1.bold = True
                current_skill_category_p = p
            else:
                if current_skill_category_p is not None:
                    if current_skill_category_p.text.endswith(': '):
                        current_skill_category_p.add_run(clean_text.strip())
                    else:
                        current_skill_category_p.add_run(", " + clean_text.strip())
                    paragraphs_to_delete.append(p)
                else:
                    p.clear()
                    p.add_run(clean_text)
                
    for p in paragraphs_to_delete:
        p._element.getparent().remove(p._element)
        
    doc.save(docx_path)
