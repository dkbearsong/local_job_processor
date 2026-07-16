import os
from datetime import datetime
from docx import Document

def _convert_markdown_to_docx(doc, markdown_text: str):
    """
    Converts a markdown string into docx elements by parsing simple markdown patterns.
    """
    if not markdown_text:
        print("No markdown text.")
        return

    print(f"Converting markdown: {markdown_text[:100]}...")
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            continue
            
        # Handle bold text (**text**)
        if '**' in line:
            # Add paragraph with inline formatting
            paragraph = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold parts (between **)
                    run = paragraph.add_run(part)
                    run.bold = True
                else:  # Regular text
                    paragraph.add_run(part)
        else:
            # Simple paragraph
            doc.add_paragraph(line)
    
    print("Successfully added to docx.")

def generate_job_report(job_title: str, company_name: str, skills_text: str, projects_text: str, matches: str, adjustments: str):
    """
    Generates a .docx report containing job analysis and suggested projects.
    Saves it to an 'output' folder with a timestamped filename.
    """
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    today_str = datetime.now().strftime("%Y-%m-%d")
    # Clean job title to remove characters that might be invalid for filenames
    if job_title != "":
        clean_job_title = "".join([c for c in job_title if c.isalnum() or c in (' ', '-', '_')]).strip()
        filename = f"{today_str}_{clean_job_title}_title report.docx"
    elif company_name != "":
        clean_company_name = "".join([c for c in company_name if c.isalnum() or c in (' ', '-', '_')]).strip()
        filename = f"{today_str}_{clean_company_name}_company report.docx"
    else:
        filename = f"{today_str}_general report.docx"
    
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    
    doc.add_heading(f'Job Analysis Report for {f"Job Title:  '{job_title}'" if job_title != '' else f"Company Name: '{company_name}'" if company_name != '' else ""}', 0)

    doc.add_heading('Skills, Responsibilities, and Requirements', level=1)
    _convert_markdown_to_docx(doc, skills_text)

    doc.add_heading('Suggested Learning Projects', level=1)
    _convert_markdown_to_docx(doc, projects_text)

    doc.add_heading('Personalized Career Matches', level=1)
    _convert_markdown_to_docx(doc, matches)

    doc.add_heading('Resume Adjustments', level=1)
    _convert_markdown_to_docx(doc, adjustments)

    try:
        doc.save(filepath)
        print(f"Successfully generated report: {filepath}")
        return filename
    except Exception as e:
        print(f"Failed to save report: {e}")
        return None